from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageSequence


DOWNLOADS = Path.home() / "Downloads"
SOURCE_NAME = "【提案评估】《斗篮机》篮球投篮机双队PK.docx"
OUTPUT_NAME = "【提案评估】《斗篮机》篮球投篮机双队PK-最新可提交版.docx"
SRC = DOWNLOADS / SOURCE_NAME
OUT = DOWNLOADS / OUTPUT_NAME

ASSET_ROOT = Path(r"E:\codexSpace\douyin-basketball-pk\logs\proposal-assets-latest")
SCREEN_DIR = ASSET_ROOT / "screenshots"
GIF_DIR = ASSET_ROOT / "gifs"

ASSETS = {
    "host": SCREEN_DIR / "01-host-config.png",
    "main": SCREEN_DIR / "02-main-ui.png",
    "gift_high_png": SCREEN_DIR / "03-high-gift.png",
    "settlement": SCREEN_DIR / "04-settlement.png",
    "red_join": GIF_DIR / "comment-red-join.gif",
    "blue_join": GIF_DIR / "comment-blue-join.gif",
    "like": GIF_DIR / "like-shot.gif",
    "gift_low": GIF_DIR / "gift-low.gif",
    "gift_mid": GIF_DIR / "gift-mid.gif",
    "gift_high": GIF_DIR / "gift-high.gif",
}

PREVIEW_DIR = ASSET_ROOT / "gif-previews"
PREVIEW_ASSETS = {
    "red_join": PREVIEW_DIR / "comment-red-join-strip.png",
    "blue_join": PREVIEW_DIR / "comment-blue-join-strip.png",
    "like": PREVIEW_DIR / "like-shot-strip.png",
    "gift_low": PREVIEW_DIR / "gift-low-strip.png",
    "gift_mid": PREVIEW_DIR / "gift-mid-strip.png",
    "gift_high": PREVIEW_DIR / "gift-high-strip.png",
}
GIF_SUMMARY_SHEET = PREVIEW_DIR / "all-gif-effects-sheet.png"


def set_run_font(run, size=None, bold=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def clear_cell(cell):
    cell._tc.clear_content()
    return cell.add_paragraph()


def clear_paragraph(paragraph):
    p = paragraph._p
    for child in list(p):
        p.remove(child)


def add_text(paragraph, text, size=10.5, bold=False):
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return run


def add_picture(paragraph, image_path, width):
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)
    return run


def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_run_font(run)


def make_row_auto(row):
    row.height = None
    row.height_rule = WD_ROW_HEIGHT_RULE.AUTO
    tr_pr = row._tr.get_or_add_trPr()
    for child in list(tr_pr):
        if child.tag == qn("w:trHeight"):
            tr_pr.remove(child)


def make_table_rows_auto(table):
    for row in table.rows:
        make_row_auto(row)


def table(doc, idx):
    if idx >= len(doc.tables):
        raise IndexError(f"Document table {idx} not found; has {len(doc.tables)} tables")
    return doc.tables[idx]


def insert_paragraph_after(anchor_element, parent):
    paragraph_element = OxmlElement("w:p")
    anchor_element.addnext(paragraph_element)
    return Paragraph(paragraph_element, parent), paragraph_element


def build_gif_preview(gif_path, out_path):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with Image.open(gif_path) as source:
        frames = [frame.convert("RGBA") for frame in ImageSequence.Iterator(source)]
    if not frames:
        raise ValueError(f"No frames in {gif_path}")

    indexes = sorted(set([0, len(frames) // 2, len(frames) - 1]))
    selected = [frames[i] for i in indexes]
    target_w = 280
    padding = 10
    label_h = 30
    resized = []
    for frame in selected:
        ratio = target_w / frame.width
        resized.append(frame.resize((target_w, int(frame.height * ratio)), Image.LANCZOS))

    max_h = max(frame.height for frame in resized)
    sheet = Image.new("RGB", (len(resized) * target_w + (len(resized) + 1) * padding, max_h + label_h + padding * 2), "#07101f")
    draw = ImageDraw.Draw(sheet)
    labels = ["START", "MID", "GOAL"]
    for idx, frame in enumerate(resized):
        x = padding + idx * (target_w + padding)
        y = padding + label_h
        draw.rounded_rectangle((x - 2, y - 2, x + target_w + 2, y + max_h + 2), radius=8, outline="#21d4fd", width=2)
        draw.text((x + 6, padding + 4), labels[idx] if idx < len(labels) else f"帧{idx + 1}", fill="#ffffff")
        bg = Image.new("RGB", (target_w, max_h), "#07101f")
        bg.paste(frame.convert("RGB"), (0, (max_h - frame.height) // 2))
        sheet.paste(bg, (x, y))
    sheet.save(out_path)


def build_gif_summary_sheet():
    rows = [
        ("RED JOIN  comment-red-join.gif", PREVIEW_ASSETS["red_join"]),
        ("BLUE JOIN  comment-blue-join.gif", PREVIEW_ASSETS["blue_join"]),
        ("LIKE SHOT  like-shot.gif", PREVIEW_ASSETS["like"]),
        ("LOW GIFT  gift-low.gif", PREVIEW_ASSETS["gift_low"]),
        ("MID GIFT  gift-mid.gif", PREVIEW_ASSETS["gift_mid"]),
        ("HIGH GIFT  gift-high.gif", PREVIEW_ASSETS["gift_high"]),
    ]
    loaded = [(label, Image.open(path).convert("RGB")) for label, path in rows]
    cols = 2
    tile_w = 470
    label_h = 34
    padding = 16
    resized = []
    for label, image in loaded:
        ratio = tile_w / image.width
        resized.append((label, image.resize((tile_w, int(image.height * ratio)), Image.LANCZOS)))

    tile_h = max(image.height for _, image in resized) + label_h + padding
    total_w = cols * tile_w + (cols + 1) * padding
    total_h = ((len(resized) + cols - 1) // cols) * tile_h + padding
    sheet = Image.new("RGB", (total_w, total_h), "#07101f")
    draw = ImageDraw.Draw(sheet)
    for idx, (label, image) in enumerate(resized):
        row = idx // cols
        col = idx % cols
        x = padding + col * (tile_w + padding)
        y = padding + row * tile_h
        draw.rounded_rectangle((x, y, x + tile_w, y + label_h - 4), radius=12, fill="#111a2e", outline="#21d4fd", width=2)
        draw.text((x + 12, y + 8), label, fill="#ffffff")
        sheet.paste(image, (x, y + label_h))
    sheet.save(GIF_SUMMARY_SHEET)


def require_inputs():
    if not SRC.exists():
        candidates = sorted(DOWNLOADS.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True)
        names = "\n".join(f"- {p.name}" for p in candidates[:20])
        raise FileNotFoundError(f"Source docx not found: {SRC}\nRecent docx files:\n{names}")
    missing = [str(p) for p in ASSETS.values() if not p.exists()]
    if missing:
        raise FileNotFoundError("Missing proposal assets:\n" + "\n".join(missing))
    build_gif_preview(ASSETS["red_join"], PREVIEW_ASSETS["red_join"])
    build_gif_preview(ASSETS["blue_join"], PREVIEW_ASSETS["blue_join"])
    build_gif_preview(ASSETS["like"], PREVIEW_ASSETS["like"])
    build_gif_preview(ASSETS["gift_low"], PREVIEW_ASSETS["gift_low"])
    build_gif_preview(ASSETS["gift_mid"], PREVIEW_ASSETS["gift_mid"])
    build_gif_preview(ASSETS["gift_high"], PREVIEW_ASSETS["gift_high"])
    build_gif_summary_sheet()


def update_document():
    require_inputs()
    doc = Document(str(SRC))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_run_font(run)
    for existing_table in doc.tables:
        format_table(existing_table)

    if len(doc.paragraphs) > 2:
        p = doc.paragraphs[2]
        clear_paragraph(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "《斗篮机》篮球投篮机双队PK", size=18, bold=True)

    try:
        checklist = table(doc, 2)
        for row_idx in [1, 2, 3, 4, 5, 6, 7, 9, 11, 12]:
            if row_idx < len(checklist.rows) and len(checklist.rows[row_idx].cells) >= 3:
                p = clear_cell(checklist.rows[row_idx].cells[2])
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_text(p, "√", size=11, bold=True)
        format_table(checklist)
    except Exception as exc:
        print(f"Checklist update skipped: {exc}")

    try:
        record = table(doc, 3)
        row = record.add_row()
        row.cells[0].text = "2026/05/12"
        row.cells[1].text = "按最新程序版本同步提案素材，确保主界面、结算页、礼物/点赞/入队动效与当前Demo一致。"
        row.cells[2].text = (
            "更新主界面交互稿、比赛结算页截图、点赞投篮GIF、红蓝入队GIF、低/中/高价礼物投篮GIF；"
            "同步最新篮板能量梁、清理篮网多余弧圈、MVP与贡献榜TOP5、礼物固定1/2/3分及合规说明。"
        )
        format_table(record)
    except Exception as exc:
        print(f"Record update skipped: {exc}")

    try:
        topic = table(doc, 5)
        if len(topic.rows) > 2 and len(topic.rows[2].cells) > 1:
            p = clear_cell(topic.rows[2].cells[1])
            add_text(
                p,
                "《斗篮机》是一款面向抖音直播间的篮球投篮机双队PK互动玩法。观众通过弹幕口令选择红队/蓝队后参与本局比赛，入队后本局不可换队；"
                "点赞与弹幕小玩法礼物均可触发投篮表现。主播可在开局前配置对局时长（3/5/10/15/20/25/30分钟）及队伍口令；点赞规则固定为30%命中率、命中得1分。"
                "低价礼物命中得1分，中价礼物命中得2分，高价礼物命中得3分；不同分档使用不同篮球皮肤、轨迹和进球特效，比赛结束后自动结算胜负、MVP和双方贡献榜TOP5。",
                size=10,
            )
    except Exception as exc:
        print(f"Topic update skipped: {exc}")

    try:
        demo = table(doc, 7)
        if len(demo.rows) > 1:
            demo.rows[1].cells[1].text = (
                "[douba-evaluation-submit.mp4]（随提案附件提交；文档内已同步插入主界面截图、结算页截图及核心指令GIF动效）"
            )
            if len(demo.rows[1].cells) > 2:
                demo.rows[1].cells[2].text = (
                    "视频覆盖主播配置、观众红蓝入队、点赞投篮、低/中/高价礼物投篮、比赛结束自动结算、MVP与贡献榜展示。"
                )
        format_table(demo)
    except Exception as exc:
        print(f"Demo update skipped: {exc}")

    try:
        framework = table(doc, 8)
        make_table_rows_auto(framework)
        cell = framework.rows[3].cells[0]
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "主界面交互稿（最新程序截图）", size=11, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture(p, ASSETS["main"], Inches(5.7))
        cap = cell.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(
            cap,
            "图：当前Demo主界面，包含红蓝队阵营卡、计时比分、篮球场、篮板篮筐、实时流水、贡献榜前三、MVP信息与主播配置面板。",
            size=9,
        )
        desc = cell.add_paragraph()
        add_text(
            desc,
            "主线流程：主播创建对局并设置时长/点赞命中率/礼物分档；观众发送“红队”或“蓝队”入队；点赞或弹幕小玩法礼物触发投篮；"
            "系统按阵营累计比分并播放对应投篮、轨迹、进球特效；时间到后自动进入结算页，展示胜负、双方得分、MVP与每队贡献榜TOP5。",
            size=10,
        )
        anchor = framework._tbl
        parent = framework._parent
        p, anchor = insert_paragraph_after(anchor, parent)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "主界面交互稿补充截图（最新程序）", size=12, bold=True)
        p, anchor = insert_paragraph_after(anchor, parent)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture(p, ASSETS["main"], Inches(6.3))
        p, anchor = insert_paragraph_after(anchor, parent)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "图：最新Demo主界面，已同步最新篮板、篮筐、篮球、投篮轨迹、进球氛围与底部信息区样式。", size=9)
    except Exception as exc:
        print(f"Framework update skipped: {exc}")

    try:
        instr = table(doc, 9)
        make_table_rows_auto(instr)
        updates = {
            1: ("发送“红队”口令，观众加入红队，本局锁定红队阵营，不可切换。", PREVIEW_ASSETS["red_join"], "动效预览：红队入队；GIF文件：comment-red-join.gif"),
            2: ("发送“蓝队”口令，观众加入蓝队，本局锁定蓝队阵营，不可切换。", PREVIEW_ASSETS["blue_join"], "动效预览：蓝队入队；GIF文件：comment-blue-join.gif"),
            3: ("其他评论仅作为普通弹幕展示，不触发入队、投篮或计分。", None, "无核心玩法动效，不参与计分"),
            4: ("低价弹幕小玩法礼物：触发一次1分投篮，命中后播放基础篮球皮肤、短轨迹与进球反馈。", PREVIEW_ASSETS["gift_low"], "动效预览：低价礼物投篮；GIF文件：gift-low.gif"),
            5: ("中价弹幕小玩法礼物：触发一次2分投篮，播放约3秒强化轨迹、能量拖尾与进球特效。", PREVIEW_ASSETS["gift_mid"], "动效预览：中价礼物投篮；GIF文件：gift-mid.gif"),
            6: ("高价弹幕小玩法礼物：触发一次3分投篮，播放约4秒高级篮球皮肤、长轨迹、篮筐能量爆发与全场氛围特效。", PREVIEW_ASSETS["gift_high"], "动效预览：高价礼物投篮；GIF文件：gift-high.gif"),
            7: ("点赞触发免费投篮，系统固定按30%命中率判定；命中得1分，用于增强直播间低门槛参与感。", PREVIEW_ASSETS["like"], "动效预览：点赞投篮；GIF文件：like-shot.gif"),
        }
        for row_idx, (desc, image_path, caption) in updates.items():
            if row_idx >= len(instr.rows):
                continue
            if len(instr.rows[row_idx].cells) > 2:
                p = clear_cell(instr.rows[row_idx].cells[2])
                add_text(p, desc, size=9.5)
            if len(instr.rows[row_idx].cells) > 3:
                cell = instr.rows[row_idx].cells[3]
                p = clear_cell(cell)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if image_path is not None:
                    add_picture(p, image_path, Inches(2.35))
                    cp = cell.add_paragraph()
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    add_text(cp, caption, size=8.5)
                else:
                    add_text(p, caption, size=9)
        format_table(instr)
        anchor = instr._tbl
        parent = instr._parent
        p, anchor = insert_paragraph_after(anchor, parent)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(p, "指令与礼物动效GIF预览（最新程序）", size=12, bold=True)
        p, anchor = insert_paragraph_after(anchor, parent)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture(p, GIF_SUMMARY_SHEET, Inches(6.3))
        p, anchor = insert_paragraph_after(anchor, parent)
        add_text(p, "说明：上图为每个GIF的“开始/过程/命中”三帧预览；原始GIF动图文件已生成，建议与提案文档、Demo视频一起上传。", size=9)
    except Exception as exc:
        print(f"Instruction update skipped: {exc}")

    try:
        design = table(doc, 10)
        make_table_rows_auto(design)
        cell = design.rows[1].cells[1]
        p = clear_cell(cell)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture(p, ASSETS["main"], Inches(3.2))
        cap1 = cell.add_paragraph()
        cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(cap1, "图1：主界面交互稿（最新Demo截图）", size=8.5)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_picture(p2, ASSETS["settlement"], Inches(3.2))
        cap2 = cell.add_paragraph()
        cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_text(cap2, "图2：比赛结束结算页（胜负、MVP、每队贡献榜TOP5）", size=8.5)
        desc = cell.add_paragraph()
        add_text(
            desc,
            "页面风格采用霓虹篮球场视觉：中央球场承载投篮动画，左右阵营卡展示队伍状态，底部信息条展示实时进球、贡献榜前三与本场MVP；"
            "结算页集中展示双方得分、获胜方、MVP及两队贡献榜TOP5，便于主播赛后复盘。",
            size=9,
        )
        format_table(design)
    except Exception as exc:
        print(f"Design screenshot update skipped: {exc}")

    try:
        settlement_table = table(doc, 14)
        if len(settlement_table.rows) > 1 and len(settlement_table.rows[1].cells) > 1:
            p = clear_cell(settlement_table.rows[1].cells[1])
            add_text(
                p,
                "对局倒计时结束后，系统自动锁定本局得分并进入结算页。结算页展示红队/蓝队最终比分、获胜阵营、本场MVP（头像/昵称/得分/命中次数）以及每队贡献榜TOP5。"
                "下一局重新开局后，观众需要重新选择阵营，避免跨局阵营状态影响公平性。",
                size=10,
            )
        if len(settlement_table.rows) > 2 and len(settlement_table.rows[2].cells) > 1:
            p = clear_cell(settlement_table.rows[2].cells[1])
            add_text(
                p,
                "MVP依据观众本局贡献分、命中次数综合展示；贡献榜按队伍分别排序，只展示前5名，防止名单过长挤压直播画面。赛后主播可根据MVP和贡献榜进行口播复盘和下一局引导。",
                size=10,
            )
        format_table(settlement_table)
    except Exception as exc:
        print(f"Settlement text update skipped: {exc}")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(p, "比赛结束结算页截图（最新程序）", size=12, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_picture(p, ASSETS["settlement"], Inches(6.3))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(cap, "图：时间截止后自动结算，展示胜负、双方最终得分、MVP与红蓝双方贡献榜TOP5。", size=9)

    p = doc.add_paragraph()
    add_text(
        p,
        "附件说明：本提案配套提交Demo演示视频 douba-evaluation-submit.mp4；文档内“指令设计”已嵌入红蓝入队、点赞投篮、低价礼物、中价礼物、高价礼物等核心GIF动效，均来自当前最新程序录制。",
        size=9,
    )

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    print(update_document())
